"""Text extraction and OCR routing.

The pipeline's job here is to turn bytes into text plus enough metadata to make
downstream decisions. The interesting decision is **when to OCR**.

OCR is expensive (seconds per page) and lossy. Most business PDFs carry a real text
layer and need none. Some are pure scans and need it for every page. Many are
mixed — a generated invoice with a scanned signature page. So the routing decision
is made **per page** on a text-density heuristic rather than per document, which
avoids both failure modes: OCR-ing a clean 40-page PDF, and returning three
characters of text for a scan because page one happened to have a header.

The threshold (`ocr_chars_per_page_threshold`, default 120) is a configurable
guess, not a measured optimum. It is documented as such in `docs/AI.md`.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

import structlog

from docflow.config import ProcessingSettings
from docflow.domain.errors import (
    CorruptDocumentError,
    EmptyDocumentError,
    OCRUnavailableError,
    TextExtractionError,
    TooManyPagesError,
)

logger = structlog.get_logger(__name__)


@dataclass(slots=True)
class PageText:
    number: int
    text: str
    used_ocr: bool = False
    char_count: int = 0

    def __post_init__(self) -> None:
        self.char_count = len(self.text.strip())


@dataclass(slots=True)
class ExtractedText:
    text: str
    pages: list[PageText] = field(default_factory=list)
    page_count: int = 0
    used_ocr: bool = False
    language: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.text.strip())

    @property
    def ocr_page_count(self) -> int:
        return sum(1 for p in self.pages if p.used_ocr)


# Ligatures and typographic characters that PDF text layers emit and that break
# naive string matching downstream (grounding, label anchoring).
_REPLACEMENTS = {
    "ﬀ": "ff",
    "ﬁ": "fi",
    "ﬂ": "fl",
    "ﬃ": "ffi",
    "ﬄ": "ffl",
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "–": "-",
    "—": "-",
    "−": "-",
    " ": " ",
    " ": " ",
    " ": " ",
    "\u200b": "",
}
_MULTI_BLANK = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+\n")


def clean_text(text: str) -> str:
    """Normalise extracted text without destroying layout.

    Layout is load-bearing: the baseline extractor anchors on labels and reads the
    remainder of the *line*, and column alignment is often the only thing
    distinguishing a value from its neighbour. So runs of spaces are preserved and
    only genuinely noisy artefacts are removed.
    """
    for source, target in _REPLACEMENTS.items():
        text = text.replace(source, target)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _TRAILING_WS.sub("\n", text)
    return _MULTI_BLANK.sub("\n\n", text).strip()


class TextExtractor:
    def __init__(self, settings: ProcessingSettings, *, max_pages: int = 50) -> None:
        self._settings = settings
        self._max_pages = max_pages

    def extract(self, data: bytes, content_type: str) -> ExtractedText:
        if content_type == "application/pdf":
            return self._from_pdf(data)
        if content_type.startswith("image/"):
            return self._from_image(data)
        if content_type == "text/plain":
            return self._from_plain(data)
        if content_type.endswith("wordprocessingml.document"):
            return self._from_docx(data)
        raise TextExtractionError(f"No text extractor for {content_type}")

    # ---------------------------------------------------------------------- pdf

    def _from_pdf(self, data: bytes) -> ExtractedText:
        import pdfplumber

        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                if len(pdf.pages) > self._max_pages:
                    raise TooManyPagesError(
                        f"Document has {len(pdf.pages)} pages; the limit is {self._max_pages}",
                        detail={"page_count": len(pdf.pages), "limit": self._max_pages},
                    )
                pages = [
                    PageText(number=i + 1, text=clean_text(page.extract_text() or ""))
                    for i, page in enumerate(pdf.pages)
                ]
                metadata = _pdf_metadata(pdf)
        except TooManyPagesError:
            raise
        except Exception as exc:
            raise CorruptDocumentError("The PDF could not be parsed for text") from exc

        sparse = [p for p in pages if p.char_count < self._settings.ocr_chars_per_page_threshold]
        if sparse and self._settings.ocr_enabled:
            self._ocr_pages(data, sparse)

        text = "\n\n".join(p.text for p in pages if p.text.strip())
        if not text.strip():
            raise EmptyDocumentError(
                "No text could be extracted from this document. "
                "If it is a scan, OCR may be unavailable or the image quality too low."
            )

        return ExtractedText(
            text=text,
            pages=pages,
            page_count=len(pages),
            used_ocr=any(p.used_ocr for p in pages),
            language=detect_language(text),
            metadata=metadata,
        )

    def _ocr_pages(self, data: bytes, pages: list[PageText]) -> None:
        """OCR the pages whose native text layer was too sparse to be real.

        Failure here is logged and swallowed rather than raised: a document with a
        good text layer on 39 of 40 pages should not fail because page 12 is a
        scanned signature that OCR could not read. The consequence — less text —
        shows up downstream as lower confidence and a review flag, which is the
        right outcome.
        """
        try:
            engine = _load_ocr_engine()
        except OCRUnavailableError as exc:
            logger.warning("ocr.unavailable", reason=exc.message, pages=len(pages))
            return

        try:
            images = engine["convert"](
                data,
                dpi=self._settings.ocr_dpi,
                first_page=min(p.number for p in pages),
                last_page=max(p.number for p in pages),
            )
        except Exception as exc:
            logger.warning("ocr.rasterisation_failed", error=type(exc).__name__)
            return

        offset = min(p.number for p in pages)
        by_number = {p.number: p for p in pages}
        for index, image in enumerate(images):
            page = by_number.get(offset + index)
            if page is None:
                continue
            try:
                recognised = engine["ocr"](image, lang=self._settings.ocr_language)
            except Exception as exc:
                logger.warning("ocr.page_failed", page=page.number, error=type(exc).__name__)
                continue
            cleaned = clean_text(recognised or "")
            # Only accept OCR output if it beats what was already there. OCR of a
            # blank page returns noise, and replacing good text with noise is worse
            # than doing nothing.
            if len(cleaned.strip()) > page.char_count:
                page.text = cleaned
                page.char_count = len(cleaned.strip())
                page.used_ocr = True

        logger.info("ocr.completed", pages=len([p for p in pages if p.used_ocr]))

    # -------------------------------------------------------------------- image

    def _from_image(self, data: bytes) -> ExtractedText:
        try:
            engine = _load_ocr_engine()
        except OCRUnavailableError:
            raise OCRUnavailableError(
                "This image needs OCR, but no OCR engine is installed. "
                "Install Tesseract, or upload a PDF with a text layer."
            ) from None

        from PIL import Image

        try:
            image = Image.open(io.BytesIO(data))
            image.load()
        except Exception as exc:
            raise CorruptDocumentError("The image could not be opened") from exc

        try:
            raw = engine["ocr"](image, lang=self._settings.ocr_language)
        except Exception as exc:
            raise TextExtractionError("OCR failed on this image") from exc

        text = clean_text(raw or "")
        if not text:
            raise EmptyDocumentError("No text could be recognised in this image")

        page = PageText(number=1, text=text, used_ocr=True)
        return ExtractedText(
            text=text,
            pages=[page],
            page_count=1,
            used_ocr=True,
            language=detect_language(text),
            metadata={"width": image.width, "height": image.height, "format": image.format},
        )

    # --------------------------------------------------------------- text/docx

    def _from_plain(self, data: bytes) -> ExtractedText:
        for encoding in ("utf-8", "utf-8-sig", "cp1250", "iso-8859-2", "latin-1"):
            try:
                raw = data.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise CorruptDocumentError("The text file uses an unsupported encoding")

        text = clean_text(raw)
        if not text:
            raise EmptyDocumentError("The text file is empty")
        return ExtractedText(
            text=text,
            pages=[PageText(number=1, text=text)],
            page_count=1,
            language=detect_language(text),
        )

    def _from_docx(self, data: bytes) -> ExtractedText:
        import docx

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:
            raise CorruptDocumentError("The Word document could not be opened") from exc

        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Tables carry most of the structured content in business documents and are
        # not in `paragraphs`. Tab-joining preserves the column relationship that
        # the label extractor depends on.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))

        text = clean_text("\n".join(parts))
        if not text:
            raise EmptyDocumentError("The Word document contains no text")
        return ExtractedText(
            text=text,
            pages=[PageText(number=1, text=text)],
            page_count=1,
            language=detect_language(text),
            metadata={"paragraphs": len(document.paragraphs), "tables": len(document.tables)},
        )


# ------------------------------------------------------------------ ocr loading

_ocr_cache: dict[str, Any] | None = None


def _load_ocr_engine() -> dict[str, Any]:
    """Resolve the OCR toolchain once, or explain precisely why it is missing.

    OCR needs both a Python binding and a native binary, and either can be absent.
    Distinguishing the two matters because the fixes are different, and a vague
    "OCR unavailable" sends the operator hunting in the wrong place.
    """
    global _ocr_cache
    if _ocr_cache is not None:
        return _ocr_cache

    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError as exc:
        raise OCRUnavailableError(
            "OCR extras are not installed. Install with: uv pip install -e '.[ocr]'"
        ) from exc

    try:
        pytesseract.get_tesseract_version()
    except Exception as exc:
        raise OCRUnavailableError(
            "The Tesseract binary was not found. Install it (macOS: `brew install tesseract`, "
            "Debian: `apt-get install tesseract-ocr`). The Docker image includes it."
        ) from exc

    _ocr_cache = {
        "ocr": lambda image, lang: pytesseract.image_to_string(image, lang=lang),
        "convert": convert_from_bytes,
    }
    return _ocr_cache


def ocr_available() -> bool:
    try:
        _load_ocr_engine()
    except OCRUnavailableError:
        return False
    return True


# ------------------------------------------------------------ pdf metadata/lang


def _pdf_metadata(pdf: Any) -> dict[str, Any]:
    raw = getattr(pdf, "metadata", None) or {}
    keep = ("Title", "Author", "Producer", "Creator", "CreationDate", "ModDate")
    return {key.lower(): str(raw[key])[:200] for key in keep if raw.get(key)}


# Stop-word frequency language detection. A full language-identification library
# would be more accurate, but language is only used to pick an OCR model and to
# label the document — a wrong guess costs nothing important, and the dependency
# would.
_LANGUAGE_MARKERS: dict[str, tuple[str, ...]] = {
    "cs": ("faktura", "dph", "částka", "splatnost", "odběratel", "dodavatel", "celkem", "účet"),
    "en": ("invoice", "amount", "payment", "total", "due", "supplier", "customer", "account"),
    "de": ("rechnung", "betrag", "zahlung", "gesamt", "fällig", "lieferant", "kunde"),
    "sk": ("faktúra", "suma", "úhrada", "dodávateľ", "odberateľ"),
    "pl": ("faktura", "kwota", "płatność", "razem", "sprzedawca", "nabywca"),
}


def detect_language(text: str) -> str | None:
    sample = text[:4000].lower()
    scores = {
        code: sum(1 for marker in markers if marker in sample)
        for code, markers in _LANGUAGE_MARKERS.items()
    }
    best = max(scores, key=lambda k: scores[k])
    return best if scores[best] >= 2 else None
